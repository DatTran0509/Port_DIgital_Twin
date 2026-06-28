# -*- coding: utf-8 -*-
"""Sinh 8 báo cáo Word (7 tính năng + chatbot) cho dự án Bản Sao Số Cảng NDT15.
Hình minh họa (sơ đồ luồng, thẻ KPI, pipeline Digital Twin, kiến trúc chatbot)
được vẽ bằng matplotlib. Chạy: python reports/gen_reports.py
"""
import os, sys, matplotlib
try: sys.stdout.reconfigure(encoding="utf-8")
except Exception: pass
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, "img"); os.makedirs(IMG, exist_ok=True)
ASSETS = os.path.join(os.path.dirname(HERE), "assets")
plt.rcParams["font.family"] = "DejaVu Sans"

# ── Dữ liệu 7 tính năng (lấy từ features.js + nội dung biên soạn thêm) ──────────
FEATURES = [
 dict(id="01", short="BẾN", color="#34E0F0", name="Điều Phối Bến Thông Minh",
   desc="Điều phối tàu tức thì & tối ưu cập bến dựa trên dữ liệu AIS real-time.",
   pain="Tắc nghẽn luồng lạch diễn ra thường xuyên do thiếu đồng bộ dữ liệu. Hãng tàu chịu phí phạt neo chờ 30.000–50.000 USD/ngày/tàu. Điều phối thủ công gây sai số ETA 45–90 phút, tạo chuỗi trễ dây chuyền.",
   mech="Hệ thống tự động quét dữ liệu định vị AIS từ xa, đối chiếu trạng thái khai thác bến real-time qua bản đồ GIS, tự động phát lệnh điều tốc tàu (Just-in-Time) để tàu không phải neo chờ.",
   mets=[("Chuẩn JIT","87%"),("Bến Trống","4 / 7"),("Tàu Theo Dõi","12"),("Nhiên Liệu Tiết Kiệm","$42k/ngày")],
   steps=[("AIS quét tàu","Nhận tín hiệu AIS 30 giây/lần trong bán kính 50 hải lý: GPS, tốc độ, hướng, IMO, DWT, ETA."),
          ("TIS đối chiếu bến","Lớp TIS tra cứu trạng thái bến real-time: đang bốc dỡ, tiến độ %, thời gian xong, sức chứa."),
          ("GIS bản đồ số","Hiển thị luồng lạch, vùng neo, mã màu tàu (Xanh JIT OK, Vàng cần chỉnh, Đỏ trễ)."),
          ("AI phát lệnh điều tốc","Thuật toán JIT tính vận tốc tối ưu, phát lệnh điều tốc qua VHF, ghi timestamp & lý do."),
          ("Berth Allocation Grid","Hiển thị tiến độ bốc dỡ theo phút, tự phân bổ bến khi tàu trước rời đi."),
          ("Xác nhận ETA","Hãng tàu nhận ETA chính xác ±10 phút qua API/portal, cập nhật tự động.")],
   layers=["AIS","TIS","GIS","AI / JIT"],
   intro="Điều phối bến là 'nút cổ chai' đầu tiên của mọi cảng: tàu đến không đúng lúc dẫn tới neo chờ, lãng phí nhiên liệu và phát sinh phí phạt lớn. Tính năng này biến việc điều phối thủ công, rời rạc thành một quy trình tự động dựa trên dữ liệu real-time, giúp tàu đến đúng thời điểm bến sẵn sàng (Just-in-Time).",
   dt_apply="Bản sao số hợp nhất ba lớp dữ liệu: AIS (định vị tàu ngoài khơi), TIS (trạng thái khai thác từng bến) và GIS (bản đồ luồng lạch). Twin liên tục cập nhật vị trí tàu và tiến độ bến theo thời gian thực, dùng AI tính toán cửa sổ cập bến tối ưu rồi mô phỏng what-if (nếu tàu A đến sớm 1 giờ thì bến nào trống?) trước khi phát lệnh điều tốc. Nhờ đó cảng chuyển từ 'phản ứng' sang 'dự báo & chủ động'.",
   apps=["Điều tốc Just-in-Time giảm thời gian neo chờ và phí phạt cho hãng tàu.",
         "Phân bổ bến tự động theo tiến độ bốc dỡ thực tế, tối đa hóa công suất cầu bến.",
         "Cảnh báo sớm xung đột lịch tàu, mô phỏng phương án thay thế trước khi xảy ra.",
         "Cung cấp ETA chính xác cho chuỗi logistics phía sau (xe, kho, hải quan)."],
   benefits=["Đạt 87% tàu cập theo chuẩn JIT, giảm mạnh neo chờ.","Tiết kiệm ~42.000 USD nhiên liệu/ngày.","Sai số ETA giảm còn ±10 phút.","Tăng công suất khai thác cầu bến mà không cần mở rộng hạ tầng."],
   plan=[("Giai đoạn 1 — Kết nối dữ liệu","Tích hợp nguồn AIS, TIS, GIS; chuẩn hóa và đồng bộ thời gian thực vào twin."),
         ("Giai đoạn 2 — Mô hình JIT","Xây thuật toán điều tốc & phân bổ bến, hiệu chỉnh theo dữ liệu lịch sử."),
         ("Giai đoạn 3 — Mô phỏng what-if","Cho phép thử kịch bản lịch tàu trên twin trước khi phát lệnh."),
         ("Giai đoạn 4 — Triển khai & đo lường","Phát lệnh qua API/VHF, theo dõi KPI JIT, tối ưu liên tục.")]),

 dict(id="02", short="BÃI", color="#27C281", name="Bãi Container Số 3D",
   desc="Số hóa bãi container & tối ưu xếp chồng.",
   pain="Tỷ lệ đảo chuyển container rỗng vô ích chiếm tới 20% chi phí bãi do quy hoạch thủ công. Thiếu tầm nhìn tổng thể khiến container bị 'chôn vùi', không lấy được đúng lúc.",
   mech="Lớp BIM phối hợp GIS & MMIS dựng mô hình số bãi 1:1. AI quét lịch trình tàu để tự xếp container: đi trước lên trên, đi sau xuống dưới — triệt tiêu đảo chuyển vô ích.",
   mets=[("Đảo Chuyển","4.1% (-20%)"),("Sức Chứa Khả Dụng","+15%"),("Năng Lượng Cẩu","-18%"),("AI Xếp Hôm Nay","1.840")],
   steps=[("Nhập lịch tàu","TIS cung cấp danh sách container theo tàu, ngày xuất, điểm đến làm đầu vào cho AI."),
          ("AI phân tích ưu tiên","Sắp xếp theo ngày xuất, trọng lượng, điểm đến để tạo bản đồ xếp chồng tối ưu 72 giờ tới."),
          ("BIM 3D dựng model","Mô hình số 3D cập nhật real-time, mỗi vị trí có địa chỉ Block-Bay-Row-Tier."),
          ("AI phân slot tự động","Mỗi container nhập cảnh được gán vị trí tự động; container xuất sớm luôn ở tầng trên."),
          ("Yard Heatmap","Bản đồ nhiệt hiển thị công suất từng khu, cảnh báo khu quá tải (>90%)."),
          ("RTG nhận lệnh","Cẩu bãi RTG nhận lệnh di chuyển tự động, giảm phụ thuộc phán đoán thủ công.")],
   layers=["TIS","BIM","GIS","MMIS","AI"],
   intro="Bãi container là 'kho hàng động' lớn nhất của cảng. Khi quy hoạch thủ công, container hay bị xếp sai thứ tự khiến cẩu phải đảo chuyển nhiều lần — vừa tốn thời gian, vừa tốn năng lượng. Tính năng này tạo một bản sao số 3D tỷ lệ 1:1 của toàn bộ bãi và để AI quyết định vị trí xếp tối ưu.",
   dt_apply="Twin dựng mô hình BIM của từng khối bãi, định vị bằng GIS và gắn dữ liệu MMIS (tình trạng thiết bị). Mỗi container có 'địa chỉ số' Block-Bay-Row-Tier cập nhật real-time. AI đọc lịch tàu (TIS) để dự báo container nào cần lấy trước rồi xếp 'đi trước lên trên', đồng thời mô phỏng bản đồ nhiệt công suất để cân bằng tải giữa các khu — toàn bộ thử nghiệm trên twin trước khi ra lệnh cho cẩu RTG.",
   apps=["Tự động gán vị trí xếp tối ưu cho mỗi container nhập cảng.",
         "Bản đồ nhiệt cảnh báo khu quá tải, đề xuất phân luồng.",
         "Tra cứu tức thì vị trí bất kỳ container nào trong bãi.",
         "Giảm quãng đường & số lần nâng hạ của cẩu, tiết kiệm năng lượng."],
   benefits=["Đảo chuyển vô ích giảm ~20% (còn 4,1%).","Tăng 15% sức chứa khả dụng nhờ xếp tối ưu.","Giảm 18% năng lượng cẩu.","Loại bỏ tình trạng container bị chôn vùi."],
   plan=[("Giai đoạn 1 — Dựng BIM bãi","Mô hình 3D 1:1 toàn bộ khối bãi, gắn hệ địa chỉ Block-Bay-Row-Tier."),
         ("Giai đoạn 2 — Thuật toán xếp chồng","AI ưu tiên theo ngày xuất/trọng lượng/điểm đến; bản đồ nhiệt công suất."),
         ("Giai đoạn 3 — Đồng bộ thiết bị","Kết nối lệnh tự động tới cẩu RTG, đồng bộ MMIS."),
         ("Giai đoạn 4 — Tối ưu liên tục","Học từ dữ liệu vận hành để giảm tiếp tỷ lệ đảo chuyển.")]),

 dict(id="03", short="CỔNG", color="#4D8DF6", name="Hệ Thống Cổng Tự Động",
   desc="Tự động hóa cổng cảng & điều tiết giao thông.",
   pain="Kẹt xe đầu kéo nhiều km tại cổng vào giờ cao điểm (Gate Chaos). Quy trình thông quan thủ công tốn ~30 phút, gây thiệt hại dây chuyền cho logistics.",
   mech="Lớp TIS kết hợp CIM thiết lập hệ thống đặt chỗ (Truck Booking Slot). Camera AI quét biển số (ALPR) đối chiếu lệnh điện tử và mở barrier dưới 2 giây — số hóa 100%.",
   mets=[("Thông Quan","1p 48s (-90%)"),("Nhận Diện Biển Số","99.4%"),("Xe Tải Hôm Nay","1.230"),("Hồ Sơ Giấy","0")],
   steps=[("Đặt Booking Slot","Công ty xe đặt lịch trước 24 giờ; hệ thống phân bổ slot giờ tránh đỉnh cao."),
          ("Nhận QR/lệnh điện tử","Tài xế nhận mã QR chứa container ID, slot giờ, cổng vào, tuyến đường."),
          ("Camera ALPR quét","Xe vào làn tự động, camera AI nhận diện biển số 99,4% trong 0,8 giây."),
          ("AI xác thực lệnh","Đối chiếu biển số với lệnh điện tử, kiểm tra slot & trọng tải <2 giây."),
          ("Barrier mở tự động","Hợp lệ → mở & hiển thị tuyến; sai → giữ và báo lý do."),
          ("Audit Log tự động","Ghi timestamp, biển số, container, hình ảnh phục vụ kiểm toán.")],
   layers=["TIS","CIM","AI / ALPR"],
   intro="Cổng là 'cửa ngõ' nơi luồng xe đường bộ gặp luồng cảng. Thông quan thủ công gây kẹt xe kéo dài và rủi ro sai sót. Tính năng này số hóa 100% quy trình ra/vào: đặt chỗ trước, nhận diện biển số tự động và xác thực lệnh điện tử trong vài giây.",
   dt_apply="Twin mô phỏng luồng xe theo thời gian thực dựa trên lớp TIS (lệnh giao nhận) và CIM (điều tiết giao thông). Hệ đặt chỗ (booking slot) san phẳng đỉnh giờ cao điểm; camera ALPR + AI xác thực đối chiếu với lệnh điện tử. Trên twin có thể mô phỏng kịch bản lưu lượng (hoặc sự cố tấn công mạng) để kiểm thử năng lực cổng trước khi xảy ra thật.",
   apps=["Đặt chỗ xe tải trước, điều tiết lưu lượng tránh kẹt.",
         "Nhận diện & thông quan tự động dưới 2 giây.",
         "Lệnh giao nhận điện tử, loại bỏ hồ sơ giấy.",
         "Audit log đầy đủ phục vụ an ninh & kiểm toán."],
   benefits=["Thời gian thông quan giảm ~90% (còn ~1p48s).","Độ chính xác nhận diện biển số 99,4%.","Hồ sơ giấy về 0.","Chấm dứt ùn tắc xe đầu kéo tại cổng."],
   plan=[("Giai đoạn 1 — Hạ tầng nhận diện","Lắp camera ALPR, barrier tự động, kết nối lệnh điện tử."),
         ("Giai đoạn 2 — Hệ đặt chỗ","Xây app booking slot, thuật toán phân bổ giờ."),
         ("Giai đoạn 3 — Tích hợp twin","Mô phỏng lưu lượng & kịch bản sự cố cổng."),
         ("Giai đoạn 4 — Vận hành & kiểm toán","Audit log tự động, tối ưu năng lực theo dữ liệu.")]),

 dict(id="05", short="AN NINH", color="#FF5468", name="Lá Chắn An Ninh AI",
   desc="Vành đai an ninh & Drone tuần tra tự động.",
   pain="Điểm mù an ninh vùng nước cảng rộng lớn trong đêm/sương mù. Chi phí canô tuần tra thủ công đắt, không phủ sóng 360°, nguy cơ xâm nhập trái phép.",
   mech="Radar quét sóng mặt nước liên tục phát hiện mục tiêu lạ. Tự động điều Drone UAV cất cánh trong 30 giây. AI phân tích video camera nhiệt xác định mức độ đe dọa.",
   mets=[("Độ Phủ Radar","360° / 5km"),("Chi Phí Tuần Tra","-50%"),("Drone Đang Bay","3 chiếc"),("Độ Chính Xác","100%")],
   steps=[("Radar SSIS quét 24/7","Mạng radar ven biển quét 360° bán kính 5km, phát hiện vật thể >0,5 knot."),
          ("AI lọc mục tiêu lạ","Loại nhiễu, phân biệt tàu đã đăng ký (AIS whitelist) với mục tiêu lạ."),
          ("UAV cất cánh tự động","Drone nhận lệnh cất cánh trong 30 giây, bay theo tọa độ GPS mục tiêu."),
          ("Camera nhiệt AI","Drone truyền video hồng ngoại; AI phân loại mục tiêu & mức đe dọa."),
          ("Alert trung tâm","Cảnh báo tức thì kèm ảnh/video và tọa độ về màn hình trung tâm."),
          ("Phối hợp ứng phó","Trung tâm xác minh và điều lực lượng, ghi nhận toàn bộ sự kiện.")],
   layers=["SSIS","Radar","UAV","AI thị giác"],
   intro="An ninh vùng nước cảng rất khó kiểm soát bằng tuần tra thủ công, đặc biệt ban đêm và sương mù. Tính năng này tạo một vành đai an ninh chủ động: radar phát hiện, AI lọc mục tiêu lạ và drone tự động xác minh — phủ sóng 360° liên tục với chi phí thấp hơn.",
   dt_apply="Lớp SSIS hợp nhất radar, camera nhiệt và đội drone vào bản sao số. Twin đối chiếu mục tiêu radar với whitelist AIS để loại tàu hợp lệ, dùng AI thị giác phân loại đe dọa và mô phỏng đường bay drone tối ưu. Có thể diễn tập kịch bản xâm nhập trên twin để kiểm tra thời gian phản ứng trước khi tình huống thật xảy ra.",
   apps=["Phát hiện & xác minh mục tiêu lạ tự động 24/7.",
         "Điều phối drone tuần tra theo tọa độ, thay canô.",
         "Phân loại mức đe dọa bằng AI camera nhiệt.",
         "Diễn tập kịch bản an ninh (xâm nhập, tấn công) trên twin."],
   benefits=["Phủ radar 360° bán kính 5km.","Giảm ~50% chi phí tuần tra.","Phản ứng drone trong 30 giây.","Loại bỏ điểm mù ban đêm/sương mù."],
   plan=[("Giai đoạn 1 — Mạng cảm biến","Triển khai radar ven biển + camera nhiệt + trạm drone."),
         ("Giai đoạn 2 — AI nhận dạng","Huấn luyện mô hình lọc mục tiêu & phân loại đe dọa."),
         ("Giai đoạn 3 — Tự động điều phối","Liên kết radar→drone→trung tâm, quy trình cảnh báo."),
         ("Giai đoạn 4 — Diễn tập & tinh chỉnh","Mô phỏng kịch bản trên twin, tối ưu thời gian phản ứng.")]),

 dict(id="08", short="MÔI TRƯỜNG", color="#2ADA9A", name="Hệ Sinh Thái Cảng Xanh (EIS)",
   desc="Quan trắc môi trường & Kiểm toán ESG tự động.",
   pain="Áp lực thuế carbon IMO nghiêm khắc; nguy cơ bị hãng tàu tẩy chay nếu không đạt ESG. Cảng thiếu dữ liệu phát thải chính xác và hệ thống báo cáo ESG tự động.",
   mech="Lớp EIS tích hợp sensor IoT đo real-time CO₂, SOx, NOx từ tàu/thiết bị và chất lượng nước biển; tổng hợp dashboard ESG và xuất báo cáo kiểm toán 1 click.",
   mets=[("Giảm CO₂","90% (khi cập bến)"),("SOₓ","0.08% ✓"),("Trạm Điện Bờ","3 kích hoạt"),("Điểm ESG Q2","84/100")],
   steps=[("IoT sensor khí thải","Trạm đo CO₂, SOx, NOx dọc cầu bến, đo 1 phút/lần."),
          ("IoT sensor nước biển","8 trạm quan trắc WQI, độ đục, pH, hydrocarbon real-time."),
          ("EIS tổng hợp","Thu nhận dữ liệu, tính chỉ số ESG, hiển thị dashboard trực quan."),
          ("Đối chiếu compliance","So sánh từng chỉ tiêu với ngưỡng IMO 2020, tạo báo cáo lỗi khi vượt."),
          ("Xuất báo cáo ESG","Một click xuất PDF/Excel chuẩn IMO cho kiểm toán độc lập."),
          ("Partner portal","Hãng tàu xem chỉ số ESG real-time để ưu tiên chọn cảng.")],
   layers=["EIS","IoT","Điện bờ","Phân tích ESG"],
   intro="Áp lực khử carbon và ESG ngày càng lớn: hãng tàu và nhà đầu tư ưu tiên cảng 'xanh' và minh bạch. Tính năng này biến môi trường từ thứ 'khó đo' thành dữ liệu real-time, tự động hóa quan trắc và báo cáo ESG.",
   dt_apply="Lớp EIS gắn cảm biến IoT (khí thải, chất lượng nước) lên bản sao số và tính chỉ số ESG liên tục. Twin đối chiếu tự động với ngưỡng IMO, mô phỏng tác động của điện bờ (giảm ~90% CO₂ khi tàu cập) và cho phép xuất báo cáo kiểm toán chỉ bằng một thao tác — biến tuân thủ môi trường thành lợi thế cạnh tranh đo được.",
   apps=["Quan trắc CO₂/SOx/NOx và chất lượng nước real-time.",
         "Tự động đối chiếu chuẩn IMO, cảnh báo vượt ngưỡng.",
         "Xuất báo cáo ESG 1-click cho kiểm toán.",
         "Cổng đối tác cho hãng tàu xem chỉ số xanh."],
   benefits=["Giảm ~90% CO₂ khi tàu cập (điện bờ).","SOx đạt ngưỡng IMO (0,08%).","Điểm ESG 84/100.","Báo cáo kiểm toán tự động, minh bạch."],
   plan=[("Giai đoạn 1 — Mạng quan trắc","Lắp sensor khí thải dọc bến + 8 trạm nước biển."),
         ("Giai đoạn 2 — Tổng hợp EIS","Xây dashboard ESG, công thức chỉ số & ngưỡng IMO."),
         ("Giai đoạn 3 — Điện bờ & giảm phát thải","Kích hoạt trạm điện bờ, mô phỏng tác động trên twin."),
         ("Giai đoạn 4 — Báo cáo & công bố","Xuất báo cáo tự động, cổng đối tác minh bạch.")]),

 dict(id="09", short="NĂNG LƯỢNG", color="#FF8A3C", name="Tài Chính & Năng Lượng",
   desc="Liên kết hoạt động vật lý với dòng tiền, mô phỏng năng lượng tái tạo.",
   pain="Chi phí năng lượng cao và khó dự báo; khó liên kết hoạt động vật lý của cảng với chi phí/dòng tiền thực tế.",
   mech="Mô phỏng sản lượng năng lượng tái tạo (gió, mặt trời) và tối ưu hóa hóa đơn tiền điện; liên kết hoạt động vật lý với dòng tiền theo thời gian thực.",
   mets=[("Chi Phí Điện","-15%"),("Điện Mặt Trời","3.2MW"),("Điện Gió","1.1MW"),("Độ Chính Xác Dự Báo","96%")],
   steps=[("IoT Sensors","Thu thập dữ liệu tiêu thụ điện theo khu vực & thiết bị."),
          ("BIM/GIS Map","Hiển thị bản đồ tiêu thụ năng lượng toàn cảng."),
          ("Power Sim","Mô phỏng sản lượng điện gió + mặt trời (microgrid)."),
          ("Price Forecast","Dự báo giá điện theo giờ, độ chính xác 96%."),
          ("Optimize Flow","Dịch tải sang giờ thấp điểm, ưu tiên nguồn tái tạo."),
          ("Cost Dashboard","Hiển thị chi phí & dòng tiền theo thời gian thực.")],
   layers=["IoT","BIM/GIS","Microgrid","AI dự báo"],
   intro="Năng lượng là chi phí lớn và biến động của cảng. Tính năng này gắn hoạt động vật lý (cẩu, chiếu sáng, điện bờ) với dòng tiền, đồng thời mô phỏng nguồn tái tạo để tối ưu hóa đơn điện và hướng tới microgrid tự chủ.",
   dt_apply="Twin thu dữ liệu tiêu thụ qua IoT, định vị trên BIM/GIS và mô phỏng sản lượng điện gió + mặt trời theo thời tiết. AI dự báo giá điện và đề xuất dịch tải sang giờ thấp điểm, ưu tiên nguồn tái tạo và pin lưu trữ (BESS). Mọi phương án tối ưu được thử trên twin trước khi áp dụng, đồng thời quy đổi tác động vật lý thành chi phí/dòng tiền real-time.",
   apps=["Mô phỏng & điều phối microgrid (gió, mặt trời, BESS).",
         "Dự báo giá điện và dịch tải tối ưu chi phí.",
         "Liên kết hoạt động vật lý với dòng tiền theo thời gian thực.",
         "Mô phỏng kịch bản mất điện và phương án nguồn dự phòng."],
   benefits=["Giảm ~15% chi phí điện.","3,2MW điện mặt trời + 1,1MW điện gió.","Độ chính xác dự báo 96%.","Tăng tự chủ năng lượng, giảm phụ thuộc lưới."],
   plan=[("Giai đoạn 1 — Đo & lập bản đồ","Lắp IoT đo tiêu thụ, dựng bản đồ năng lượng BIM/GIS."),
         ("Giai đoạn 2 — Mô phỏng nguồn","Mô hình hóa điện gió/mặt trời/BESS thành microgrid."),
         ("Giai đoạn 3 — Dự báo & tối ưu","AI dự báo giá điện, thuật toán dịch tải."),
         ("Giai đoạn 4 — Liên kết tài chính","Dashboard chi phí/dòng tiền real-time, báo cáo.")]),

 dict(id="10", short="TỔNG QUAN", color="#B07CFF", name="Tổng Quan Bến Cảng & CHRONOS",
   desc="Góc nhìn bao quát toàn cảnh cảng thông minh và cỗ máy thời gian CHRONOS.",
   pain="Lãnh đạo thiếu một 'bức tranh sống' duy nhất để giám sát toàn cảng và thiếu công cụ thử nghiệm 'điều gì xảy ra nếu…' trước khi ra quyết định lớn.",
   mech="Hợp nhất toàn bộ 15 lớp dữ liệu vào một cảnh 3D sống; bổ sung cỗ máy thời gian CHRONOS cho phép tua quá khứ, xem tương lai (precognition) và phân nhánh thực tại (fork) để mô phỏng sự cố.",
   mets=[("Tàu Đang Cập Bến","4"),("Xe Tải Đang Chạy","12"),("Cần Cẩu Hoạt Động","8"),("UAV Tuần Tra","3")],
   steps=[("Góc nhìn rộng","Quan sát toàn bộ quy mô bến cảng trong một cảnh 3D."),
          ("Theo dõi luồng","Xem sự phối hợp giữa tàu, xe, cẩu theo thời gian thực."),
          ("Tua thời gian (CHRONOS)","Tua ngược/dừng/tua nhanh; mọi thực thể phục tùng một dòng thời gian."),
          ("Xem tương lai","Kéo vượt hiện tại để thấy lớp 'bóng ma' dự báo vị trí tàu."),
          ("Fork Reality","Tiêm sự cố (cẩu hỏng, bão, mất điện…) và mô phỏng hậu quả cascade."),
          ("Glass Box","Click hậu quả để xem chuỗi nhân quả ngược về gốc.")],
   layers=["15 lớp dữ liệu","CHRONOS","Mô phỏng what-if","Glass Box"],
   intro="Tổng Quan là 'buồng lái' của lãnh đạo: một bức tranh sống hợp nhất toàn bộ 15 lớp dữ liệu. Điểm đột phá là cỗ máy thời gian CHRONOS — biến bản sao số từ 'tấm gương' phản chiếu hiện tại thành 'nhà tiên tri' có thể tua thời gian và mô phỏng các tương lai khác nhau.",
   dt_apply="Đây là nơi sức mạnh digital twin thể hiện rõ nhất: toàn bộ dữ liệu real-time hội tụ vào một cảnh 3D. CHRONOS kiểm soát 'dòng thời gian mô phỏng' — cho phép tua lại lịch sử đã ghi, nhìn trước tương lai dự báo, và đặc biệt là PHÂN NHÁNH thực tại để chạy kịch bản 'điều gì xảy ra nếu…' với hậu quả lan truyền qua bến→bãi→cổng. Đây là bước nhảy lên cấp Prescriptive (cấp 4/5) — không chỉ giám sát mà còn đề xuất hành động.",
   apps=["Giám sát toàn cảng trong một màn hình duy nhất.",
         "Tua lại sự kiện để điều tra/đào tạo.",
         "Dự báo tắc nghẽn & nhu cầu trước khi xảy ra.",
         "Mô phỏng sự cố (bão, mất điện, tấn công) và thử phương án hóa giải."],
   benefits=["Một 'bức tranh sống' duy nhất cho lãnh đạo.","Ra quyết định dựa trên mô phỏng, giảm rủi ro.","Nâng cấp độ twin lên Prescriptive.","Công cụ trình diễn mạnh cho nhà đầu tư & chính phủ."],
   plan=[("Giai đoạn 1 — Hội tụ dữ liệu","Hợp nhất 15 lớp vào một cảnh 3D real-time."),
         ("Giai đoạn 2 — Lõi thời gian","Xây CHRONOS: tua/dừng/tua nhanh + ghi snapshot."),
         ("Giai đoạn 3 — Tương lai & Fork","Lớp dự báo (ghost) + engine mô phỏng sự cố cascade."),
         ("Giai đoạn 4 — Glass Box & Copilot","Giải thích nhân quả + trợ lý hội thoại điều khiển twin.")]),

 dict(id="11", short="THIÊN TAI", color="#E0556B", name="Mô Phỏng Thiên Tai & Sự Cố (CHRONOS Resilience)",
   desc="Phân nhánh thực tại để mô phỏng thiên tai & sự cố (bão, ngập, mất điện, tấn công mạng, hỏng thiết bị) và thử phương án hóa giải trước khi xảy ra thật.",
   pain="Thiên tai và sự cố vận hành — bão, triều cường, nước biển dâng, mất điện lưới, tấn công mạng, hỏng cẩu — gây gián đoạn dây chuyền và thiệt hại rất lớn. Cảng thường thiếu công cụ DIỄN TẬP & DỰ PHÒNG trước khi sự cố xảy ra, khiến ứng phó bị động, quyết định đầu tư hạ tầng chống chịu thiếu cơ sở định lượng.",
   mech="Nút 'Fork Reality' tiêm một chuỗi sự kiện theo mốc thời gian vào bản sao số, chạy SONG SONG hai vũ trụ (Gốc và Phản thực), dùng mô hình cascade lan truyền hậu quả qua bến → bãi → cổng → chi phí/CO₂, hiển thị biểu đồ phân kỳ, rồi cho phép 'AI Hóa Giải' và truy vết nhân quả (Glass Box).",
   mets=[("Kịch bản dựng sẵn","5+"),("So sánh","Gốc vs Phản thực"),("Hóa giải","AI mitigation"),("Cấp độ Twin","Prescriptive")],
   steps=[("Chọn kịch bản","Mở bảng Fork Reality, chọn sự cố: cẩu hỏng, dồn tàu, bão+ngập, tấn công mạng, mất điện."),
          ("Tiêm sự kiện theo mốc","Hệ tiêm chuỗi sự kiện có thời điểm (vd t+8' nước +1.8m, t+20' mất điện) vào twin."),
          ("Chạy cascade song song","Mô phỏng đồng thời vũ trụ Gốc (không sự cố) và Phản thực (có sự cố) theo từng phút."),
          ("Hiển thị tác động","Biểu đồ phân kỳ + thẻ KPI (chờ tàu, kẹt cổng, chi phí, CO₂) + sơ đồ ảnh hưởng tới cảng."),
          ("AI Hóa Giải","Áp phương án đối ứng (điều cẩu, bơm thoát nước, chuyển nguồn) → đường phản thực uốn lại."),
          ("Glass Box nhân quả","Click vào hậu quả để xem chuỗi sợi sáng truy ngược về nguyên nhân gốc.")],
   layers=["CHRONOS","Mô hình cascade","EIS / An ninh","AI hóa giải"],
   intro="Đây là tính năng 'phòng thủ' chiến lược của cảng: thay vì chờ thiên tai/sự cố xảy ra rồi mới ứng phó, bản sao số cho phép DIỄN TẬP trước trong môi trường ảo. Người dùng tiêm một sự cố vào twin và xem hậu quả lan truyền khắp cảng theo thời gian, rồi thử các phương án hóa giải để chọn cách tối ưu — biến rủi ro từ 'bất ngờ' thành 'đã có kịch bản'.",
   dt_apply="Tính năng khai thác trọn vẹn sức mạnh digital twin ở cấp Prescriptive. Lõi thời gian CHRONOS cho phép phân nhánh dòng thời gian: một nhánh giữ nguyên thực tại (baseline), một nhánh chịu sự cố (forked). Mô hình cascade nhẹ lan truyền tác động qua các nút bến → bãi → cổng → năng lượng, quy đổi thành chi phí và phát thải. Hiệu ứng được trực quan hóa ngay trong cảnh 3D (nước dâng, trạm điện ngầm ngập đỏ, cẩu ngừng, drone hạ cánh) song song với biểu đồ so sánh. Glass Box ghi lại cây nhân quả để giải thích 'vì sao'.",
   apps=["Diễn tập ứng phó bão, triều cường, nước biển dâng (chống chịu khí hậu).",
         "Mô phỏng mất điện lưới và kiểm thử nguồn dự phòng/microgrid.",
         "Diễn tập tấn công mạng vào cổng và quy trình cô lập/khôi phục.",
         "Mô phỏng hỏng thiết bị (cẩu) và phương án điều chuyển nguồn lực.",
         "Định lượng thiệt hại (tiền/CO₂/thời gian) của từng kịch bản để lập kế hoạch dự phòng."],
   benefits=["Chuyển từ ứng phó bị động sang chủ động có kịch bản.","Định lượng được thiệt hại & hiệu quả hóa giải bằng số liệu.","Đào tạo kíp vận hành an toàn trong môi trường ảo.","Cơ sở định lượng cho quyết định đầu tư hạ tầng chống chịu (đê, bơm, nguồn dự phòng)."],
   plan=[("Giai đoạn 1 — Mô hình cascade","Xây mô hình lan truyền hậu quả bến→bãi→cổng→năng lượng, hiệu chỉnh theo dữ liệu lịch sử."),
         ("Giai đoạn 2 — Thư viện kịch bản","Số hóa 5 kịch bản lõi + tham số (cường độ bão, mực nước, thời lượng mất điện)."),
         ("Giai đoạn 3 — Hiệu ứng & trực quan","Liên kết hiệu ứng 3D (ngập, mất điện) + biểu đồ phân kỳ + Glass Box."),
         ("Giai đoạn 4 — Hóa giải & diễn tập","Bộ phương án AI hóa giải + quy trình diễn tập định kỳ cho kíp vận hành.")],
   scenarios=[("🛠 Cẩu B-03 hỏng","Cẩu ngừng → dwell tăng → bãi ùn ứ → cổng kẹt → chi phí & CO₂ tăng.","Điều cẩu B-02 sang hỗ trợ."),
              ("🚢 Dồn 3 tàu","3 tàu xin cập cùng cửa sổ thủy triều, thiếu bến → phí neo chờ tăng vọt.","JIT điều tốc để tàu đến đúng giờ."),
              ("🌊 Bão + nước dâng","Gió mạnh (cẩu ngừng) + nước dâng +2.4m → trạm điện ngầm ngập → cổng thủ công → cảng tê liệt.","Bơm thoát nước + chuyển nguồn dự phòng."),
              ("🛡 Tấn công mạng cổng","ALPR/booking bị khóa → cổng thủ công → throughput ÷3 → hàng xe kéo dài.","An ninh AI cô lập & khôi phục cổng."),
              ("⚡ Mất điện lưới","Mất lưới → cẩu & cổng giảm tải → nguồn tái tạo gánh → cắt tải luân phiên.","Điện mặt trời/gió + BESS gánh tải cốt lõi.")]),
]

PRIMARY = "#0B1220"

# ── Nội dung chuyên sâu bổ sung (để mỗi báo cáo đạt 6-7 trang) ─────────────────
EXTRA = {
 "01": dict(
   context="Cập bến là điểm khởi đầu của chuỗi giá trị cảng và cũng là nơi phát sinh chi phí ẩn lớn nhất. Khi đội tàu ngày càng khổng lồ (ULCV 18.000–24.000 TEU) và lịch trình toàn cầu siết chặt theo từng giờ, điều phối bến chính xác trở thành lợi thế cạnh tranh sống còn: tàu đến sai thời điểm phải neo chờ ngoài luồng, đốt nhiên liệu và chịu phí phạt, kéo theo trễ dây chuyền cho cả mạng lưới.",
   tech="Tính năng vận hành trên ba lớp dữ liệu lõi. AIS thu tín hiệu định vị tàu 30 giây/lần trong bán kính 50 hải lý (GPS, tốc độ, hướng, IMO, DWT, ETA). TIS cung cấp trạng thái khai thác từng bến theo thời gian thực. GIS dựng bản đồ luồng lạch, vùng neo và vị trí bến. Một động cơ AI (thuật toán JIT) chạy trên bản sao số hợp nhất ba lớp này, tính cửa sổ cập bến tối ưu và phát lệnh điều tốc qua VHF/API.",
   deepdive="Khác biệt cốt lõi nằm ở vòng lặp dự báo – đối chiếu – điều chỉnh. Thay vì chờ tàu đến rồi mới phân bến (phản ứng), hệ thống dự báo thời điểm bến trống dựa trên tiến độ bốc dỡ thực tế, đối chiếu với ETA của các tàu đang tới, rồi mô phỏng nhiều phương án phân bổ trên twin trước khi chọn phương án tối ưu. Mọi thay đổi (tàu đến sớm/muộn, sự cố cẩu) được tái tính tức thì để giữ lịch luôn khả thi.",
   roi="Với phí neo chờ 30.000–50.000 USD/ngày/tàu và hàng chục lượt tàu mỗi tuần, nâng tỷ lệ JIT lên 87% cắt giảm đáng kể chi phí cho hãng tàu và tăng vòng quay cầu bến. Vì chủ yếu là tối ưu phần mềm trên hạ tầng dữ liệu sẵn có nên thời gian hoàn vốn rất ngắn.",
   risks=[("Chất lượng/độ trễ dữ liệu AIS","Bổ sung AIS vệ tinh + nội suy quỹ đạo, kiểm tra tính nhất quán trước khi phát lệnh."),
          ("Hãng tàu không tuân thủ lệnh điều tốc","Chuẩn hóa giao thức JIT, chia sẻ lợi ích tiết kiệm nhiên liệu để tạo động lực."),
          ("Sai số ETA do thời tiết/luồng","Tích hợp dữ liệu thủy văn–khí tượng, cập nhật liên tục và đặt biên an toàn động.")],
   conclusion="Điều phối bến thông minh là đòn bẩy chi-phí-thấp, tác-động-cao: khai thác dữ liệu sẵn có để loại bỏ neo chờ và tăng công suất cầu bến. Khuyến nghị triển khai sớm vì đây cũng là nền dữ liệu (AIS/TIS/GIS) cho các tính năng phía sau."),
 "02": dict(
   context="Bãi container là 'kho hàng động' lớn nhất của cảng, nơi mỗi thao tác thừa đều nhân lên theo hàng nghìn container/ngày. Quy hoạch thủ công khiến container bị xếp sai thứ tự, cẩu phải đảo chuyển nhiều lần — tiêu tốn tới ~20% chi phí bãi và làm chậm cả chuỗi giao nhận.",
   tech="Lớp BIM dựng mô hình số 3D tỷ lệ 1:1 của từng khối bãi; GIS định vị; MMIS gắn tình trạng thiết bị. Mỗi container có 'địa chỉ số' Block-Bay-Row-Tier cập nhật real-time. AI đọc lịch tàu từ TIS để dự báo container nào cần lấy trước và sinh bản đồ nhiệt công suất từng khu.",
   deepdive="Thuật toán xếp chồng tối ưu theo ba tiêu chí: ngày xuất (đi trước lên trên), trọng lượng (nặng xuống dưới) và điểm đến (gom nhóm). Trước khi ra lệnh cho cẩu RTG, hệ mô phỏng kết quả trên twin và cân bằng tải giữa các khu để tránh điểm nóng >90% — toàn bộ là 'thử trước, làm sau'.",
   roi="Giảm ~20% đảo chuyển vô ích, tăng 15% sức chứa khả dụng và giảm 18% năng lượng cẩu đồng nghĩa với việc nâng năng lực khai thác mà KHÔNG cần mở rộng hạ tầng vật lý — khoản tiết kiệm trực tiếp và liên tục.",
   risks=[("Sai lệch giữa twin và bãi thực","Đồng bộ vị trí qua quét định kỳ (OCR/RFID/định vị cẩu), cảnh báo khi lệch."),
          ("Thay đổi lịch tàu đột ngột","Tái tối ưu xếp chồng theo sự kiện, giữ vùng đệm linh hoạt."),
          ("Phụ thuộc thiết bị cẩu","Theo dõi tình trạng MMIS, lập kế hoạch bảo trì dự đoán.")],
   conclusion="Số hóa bãi 3D biến 'điểm mù' lớn nhất của cảng thành không gian minh bạch, để AI loại bỏ thao tác thừa. Đây là tính năng cho hiệu quả khai thác cao và nên đi cùng tính năng Điều phối bến."),
 "03": dict(
   context="Cổng là cửa ngõ nơi luồng xe đường bộ gặp luồng cảng — và là nơi 'Gate Chaos' xảy ra: kẹt xe đầu kéo nhiều km vào giờ cao điểm, thông quan thủ công ~30 phút/lượt gây thiệt hại dây chuyền cho logistics nội địa.",
   tech="Lớp TIS (lệnh giao nhận) kết hợp CIM (điều tiết giao thông). Camera ALPR + AI nhận diện biển số 99,4% trong <1 giây, đối chiếu lệnh điện tử và điều khiển barrier; toàn bộ ghi audit log. Hệ đặt chỗ (booking slot) phân bổ giờ để san phẳng đỉnh cao điểm.",
   deepdive="Cơ chế 'đặt chỗ trước' là chìa khóa: thay vì để xe đổ dồn ngẫu nhiên, hệ phân bổ slot theo năng lực cổng và bãi đích, biến dòng xe hỗn loạn thành luồng có lịch. Twin cho phép mô phỏng lưu lượng (và cả kịch bản tấn công mạng) để kiểm thử năng lực trước khi xảy ra thật.",
   roi="Thông quan giảm ~90% (còn ~1p48s), hồ sơ giấy về 0, chấm dứt ùn tắc — vừa tiết kiệm chi phí vận hành vừa nâng trải nghiệm cho hãng vận tải, tăng sức hút thương mại của cảng.",
   risks=[("Nhận diện sai biển số (mưa, bẩn)","Mô hình ALPR đa điều kiện + xác thực phụ (QR/RFID) khi điểm tin cậy thấp."),
          ("Tấn công mạng vào hệ cổng","Phân vùng IT/OT, giám sát bất thường, làn thủ công dự phòng (xem kịch bản sự cố)."),
          ("Tài xế không đặt slot","Ưu tiên làn cho xe có booking, truyền thông & ưu đãi để tạo thói quen.")],
   conclusion="Cổng tự động là 'bộ mặt' số hóa dễ thấy nhất của cảng với ROI nhanh. Nên triển khai cùng lớp an ninh để vừa thông suốt vừa an toàn."),
 "05": dict(
   context="Vùng nước cảng rộng lớn rất khó kiểm soát bằng tuần tra thủ công, đặc biệt ban đêm và sương mù. Canô tuần tra tốn kém, không phủ 360°, để lại điểm mù — rủi ro xâm nhập trái phép và mất an ninh chuỗi cung ứng.",
   tech="Lớp SSIS hợp nhất mạng radar ven biển (quét 360°, bán kính 5km), camera nhiệt và đội drone UAV vào bản sao số. AI thị giác phân loại mục tiêu; hệ đối chiếu mục tiêu radar với whitelist AIS để loại tàu hợp lệ và chỉ tập trung vào mục tiêu lạ.",
   deepdive="Quy trình tự động khép kín: radar phát hiện → AI lọc nhiễu & whitelist → drone cất cánh trong 30 giây bay tới tọa độ → camera nhiệt phân loại đe dọa → cảnh báo trung tâm kèm ảnh/tọa độ. Twin cho phép diễn tập kịch bản xâm nhập để đo thời gian phản ứng và tối ưu đường bay drone.",
   roi="Giảm ~50% chi phí tuần tra, phủ sóng 360° liên tục và loại bỏ điểm mù ban đêm — nâng mức an ninh đạt chuẩn quốc tế (ISPS) đồng thời giảm chi phí vận hành đội canô.",
   risks=[("Báo động giả (chim, sóng)","Huấn luyện AI lọc nhiễu, kết hợp đa cảm biến để xác nhận chéo."),
          ("Điều kiện thời tiết hạn chế drone","Bố trí drone dự phòng + cảm biến cố định, hạ ngưỡng khi thời tiết xấu."),
          ("An toàn bay & pháp lý UAV","Tuân thủ vùng bay, quản lý đường bay tự động, ghi nhật ký đầy đủ.")],
   conclusion="Lá chắn an ninh AI biến an ninh thụ động thành chủ động 360°, chi phí thấp hơn. Là tính năng tạo niềm tin cho hãng tàu và cơ quan quản lý."),
 "08": dict(
   context="Áp lực khử carbon và ESG ngày càng quyết liệt: thuế/giá carbon (IMO, EU ETS) tăng, còn hãng tàu và nhà đầu tư ưu tiên cảng 'xanh', minh bạch. Cảng thiếu dữ liệu phát thải chính xác sẽ vừa chịu rủi ro chi phí vừa mất lợi thế thương mại.",
   tech="Lớp EIS tích hợp cảm biến IoT đo CO₂/SOx/NOx dọc cầu bến (1 phút/lần) và 8 trạm quan trắc nước biển (WQI, độ đục, pH, hydrocarbon). Dữ liệu được tổng hợp thành chỉ số ESG, đối chiếu tự động với ngưỡng IMO và xuất báo cáo chuẩn chỉ bằng một thao tác.",
   deepdive="Điểm mạnh là biến môi trường — vốn 'khó đo' — thành dữ liệu real-time có thể hành động. Twin mô phỏng tác động của điện bờ (giảm ~90% CO₂ khi tàu cập), cảnh báo tức thì khi một chỉ tiêu vượt ngưỡng, và cung cấp cổng đối tác để hãng tàu xem chỉ số xanh trước khi chọn cảng.",
   roi="Tránh rủi ro thuế carbon, đạt điểm ESG cao (84/100) và trở thành lựa chọn ưu tiên của các hãng tàu cam kết net-zero — biến tuân thủ môi trường thành lợi thế cạnh tranh đo được bằng doanh thu.",
   risks=[("Độ chính xác & hiệu chuẩn cảm biến","Lịch hiệu chuẩn định kỳ, đối chiếu phòng thí nghiệm, cảm biến dự phòng."),
          ("Thay đổi quy chuẩn phát thải","Thiết kế ngưỡng cấu hình được để cập nhật theo IMO/EU ETS."),
          ("Chi phí hạ tầng điện bờ","Triển khai theo giai đoạn ở bến ưu tiên, đo ROI trước khi mở rộng.")],
   conclusion="Cảng Xanh EIS vừa là 'tấm khiên' tuân thủ vừa là 'nam châm' thu hút hãng tàu xanh. Nên gắn chặt với tính năng Năng lượng để cộng hưởng hiệu quả giảm phát thải."),
 "09": dict(
   context="Năng lượng là một trong những chi phí lớn và biến động nhất của cảng, lại khó liên kết với hoạt động vật lý và dòng tiền. Khi giá điện dao động mạnh và xu hướng điện khí hóa thiết bị tăng, quản trị năng lượng thông minh trở thành yếu tố quyết định biên lợi nhuận.",
   tech="Cảm biến IoT đo tiêu thụ theo khu vực/thiết bị; BIM/GIS dựng bản đồ năng lượng; mô-đun mô phỏng nguồn tái tạo (điện gió 1,1MW + mặt trời 3,2MW + pin BESS) tạo microgrid. AI dự báo giá điện theo giờ (độ chính xác 96%) và đề xuất dịch tải.",
   deepdive="Hệ tối ưu theo hai trục: dịch tải sang giờ thấp điểm và ưu tiên nguồn tái tạo/pin lưu trữ. Mọi phương án được thử trên twin trước khi áp dụng, đồng thời quy đổi tác động vật lý (giờ chạy cẩu, chiếu sáng, điện bờ) thành chi phí/dòng tiền real-time để lãnh đạo thấy ngay ảnh hưởng tài chính.",
   roi="Giảm ~15% chi phí điện cùng việc tăng tự chủ năng lượng (giảm phụ thuộc lưới) mang lại tiết kiệm trực tiếp và lá chắn trước biến động giá — đặc biệt giá trị khi quy mô tiêu thụ lớn.",
   risks=[("Biến động nguồn tái tạo (thời tiết)","Kết hợp BESS + dự báo thời tiết, giữ công suất lưới dự phòng."),
          ("Dự báo giá điện sai lệch","Học liên tục từ dữ liệu thị trường, đặt biên an toàn cho lịch dịch tải."),
          ("Đầu tư ban đầu microgrid","Triển khai theo lộ trình, ưu tiên hạng mục hoàn vốn nhanh trước.")],
   conclusion="Tài chính & Năng lượng nối 'thế giới vật lý' với 'dòng tiền', biến tiết kiệm năng lượng thành con số trên báo cáo. Là nền cho lộ trình net-zero cùng Green Hub."),
 "10": dict(
   context="Lãnh đạo cảng cần một 'buồng lái' duy nhất để thấy toàn cảnh và một công cụ để hỏi 'điều gì xảy ra nếu…' trước khi ra quyết định lớn. Tổng Quan & CHRONOS đáp ứng cả hai: hội tụ 15 lớp dữ liệu vào một cảnh sống và bổ sung khả năng tua thời gian.",
   tech="Toàn bộ dữ liệu real-time (AIS, TIS, GIS, BIM, EIS, SSIS, năng lượng…) hội tụ vào một cảnh 3D. Lõi CHRONOS kiểm soát 'dòng thời gian mô phỏng' với cơ chế ghi snapshot để tua lại, lớp dự báo (ghost) để nhìn trước, và engine phân nhánh (fork) để mô phỏng kịch bản.",
   deepdive="CHRONOS tách thời gian mô phỏng khỏi đồng hồ tường: mọi thực thể (tàu, cẩu, xe) phục tùng một dòng thời gian duy nhất, nên có thể tua ngược (replay từ snapshot), tua tới và đặc biệt là phân nhánh để chạy 'vũ trụ phản thực'. Glass Box bổ sung khả năng giải thích nhân quả — biến hộp đen AI thành hộp kính.",
   roi="Giá trị nằm ở chất lượng quyết định: lãnh đạo ra quyết định dựa trên mô phỏng thay vì phỏng đoán, giảm rủi ro của các khoản đầu tư hạ tầng lớn. Đồng thời là công cụ trình diễn mạnh để thu hút nhà đầu tư & đối tác.",
   risks=[("Khối lượng dữ liệu & hiệu năng","Tối ưu render/lưu trữ, dùng snapshot nhẹ và mô hình cascade rời rạc."),
          ("Tin cậy của mô phỏng what-if","Nêu rõ giả định, hiệu chỉnh theo dữ liệu thật, dùng cho định hướng quyết định."),
          ("Đào tạo người dùng","Bổ sung trợ lý hội thoại (Copilot) và tham quan có dẫn để hạ rào cản sử dụng.")],
   conclusion="Tổng Quan & CHRONOS là 'vương miện' của dự án — nơi digital twin lên cấp Prescriptive. Nên là điểm nhấn khi trình bày với lãnh đạo và nhà đầu tư."),
 "11": dict(
   context="Biến đổi khí hậu (bão mạnh hơn, nước biển dâng) cùng các rủi ro vận hành & an ninh mạng khiến 'khả năng chống chịu' (resilience) trở thành tiêu chí sống còn của cảng hiện đại. Chi phí một ngày tê liệt cảng là rất lớn, nên năng lực diễn tập & dự phòng TRƯỚC sự cố là khoản đầu tư đáng giá.",
   tech="Tính năng dựng trên lõi CHRONOS (phân nhánh thời gian) + một mô hình cascade nhẹ lan truyền hậu quả qua bến→bãi→cổng→năng lượng, kết hợp dữ liệu EIS (thủy văn) và an ninh. Hiệu ứng được hiển thị đồng thời trong cảnh 3D (nước dâng, ngập trạm điện ngầm, cẩu ngừng) và biểu đồ phân kỳ Gốc–Phản thực.",
   deepdive="Khi 'Fork Reality', hệ chạy SONG SONG hai vũ trụ và quy đổi hậu quả thành chi phí/CO₂/thời gian chờ theo từng phút, cho thấy độ lớn và tốc độ lan của thiệt hại. Nút 'AI Hóa Giải' tiêm phương án đối ứng để đo hiệu quả giảm thiểu; Glass Box truy vết nhân quả về gốc — phục vụ cả ra quyết định lẫn đào tạo.",
   roi="Giá trị đến từ thiệt hại TRÁNH ĐƯỢC: diễn tập trước giúp rút ngắn thời gian phục hồi, tránh quyết định sai lúc khủng hoảng, và cung cấp cơ sở định lượng để đầu tư đúng vào hạ tầng chống chịu (đê, bơm thoát nước, nguồn dự phòng, an ninh mạng).",
   risks=[("Mô hình cascade là gần đúng","Hiệu chỉnh tham số theo sự cố lịch sử; dùng cho định hướng, không thay thế kế hoạch ƯPSC chính thức."),
          ("Kịch bản chưa bao quát hết","Cho phép tùy biến tham số & thêm kịch bản mới; rà soát định kỳ với chuyên gia."),
          ("Phụ thuộc chất lượng dữ liệu nền","Tăng độ phủ cảm biến (thủy văn, điện, an ninh) để mô phỏng sát thực tế hơn.")],
   conclusion="Mô phỏng Thiên tai & Sự cố nâng cảng từ 'ứng phó bị động' lên 'chủ động có kịch bản'. Đây là tính năng tạo khác biệt mạnh với cấp lãnh đạo và cơ quan quản lý, nên được diễn tập định kỳ."),
}
for _f in FEATURES:
    _f.update(EXTRA.get(_f["id"], {}))

# ── Vẽ hình ────────────────────────────────────────────────────────────────────
def save(fig, path):
    fig.savefig(path, dpi=130, bbox_inches="tight", facecolor="white"); plt.close(fig)

import textwrap
def wrp(t, w, maxlines=99):
    lines = textwrap.wrap(t, width=w)
    if len(lines) > maxlines:
        lines = lines[:maxlines]; lines[-1] = lines[-1].rstrip("…").rstrip() + "…"
    return "\n".join(lines)

def flow_img(f):
    steps = f["steps"]; col = f["color"]; path = os.path.join(IMG, f"flow_{f['id']}.png")
    fig, ax = plt.subplots(figsize=(11, 6.2)); ax.axis("off"); ax.set_xlim(0, 12); ax.set_ylim(0, 7)
    W, H = 3.3, 2.5
    cx = [2.05, 6.0, 9.95]; cyt, cyb = 5.15, 1.55
    cells = [(cx[0],cyt),(cx[1],cyt),(cx[2],cyt),(cx[2],cyb),(cx[1],cyb),(cx[0],cyb)]
    for i,(s,(x,y)) in enumerate(zip(steps, cells)):
        ax.add_patch(FancyBboxPatch((x-W/2, y-H/2), W, H, boxstyle="round,pad=0.02,rounding_size=0.14",
                     linewidth=1.8, edgecolor=col, facecolor=col+"16"))
        ax.text(x-W/2+0.22, y+H/2-0.30, str(i+1), fontsize=16, fontweight="bold", color=col, ha="left", va="center")
        ax.text(x, y+H/2-0.34, wrp(s[0], 26, 2), fontsize=10, fontweight="bold", color=PRIMARY, ha="center", va="top")
        ax.text(x, y+0.02, wrp(s[1], 36, 4), fontsize=7.0, color="#33415C", ha="center", va="top", linespacing=1.3)
    arrows = [((cx[0]+W/2, cyt),(cx[1]-W/2, cyt)), ((cx[1]+W/2, cyt),(cx[2]-W/2, cyt)),
              ((cx[2], cyt-H/2),(cx[2], cyb+H/2)), ((cx[2]-W/2, cyb),(cx[1]+W/2, cyb)),
              ((cx[1]-W/2, cyb),(cx[0]+W/2, cyb))]
    for a,b in arrows:
        ax.add_patch(FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=18, color="#9AA3B2", lw=1.8))
    save(fig, path); return path

def kpi_img(f):
    mets=f["mets"]; col=f["color"]; path=os.path.join(IMG,f"kpi_{f['id']}.png")
    fig, axes = plt.subplots(1, 4, figsize=(10, 2.4)); fig.subplots_adjust(wspace=0.25)
    for ax,(label,val) in zip(axes, mets):
        ax.axis("off"); ax.set_xlim(0,1); ax.set_ylim(0,1)
        ax.add_patch(FancyBboxPatch((0.06,0.1),0.88,0.8, boxstyle="round,pad=0.02,rounding_size=0.12",
                     linewidth=1.6, edgecolor=col, facecolor=col+"16"))
        ax.text(0.5, 0.64, wrp(val, 12, 2), fontsize=12.5, fontweight="bold", color=col, ha="center", va="center")
        ax.text(0.5, 0.28, wrp(label, 15, 2), fontsize=8.0, color="#33415C", ha="center", va="center", linespacing=1.2)
    save(fig,path); return path

def pipeline_img(f):
    col=f["color"]; path=os.path.join(IMG,f"pipe_{f['id']}.png")
    stages=[("NGUỒN DỮ LIỆU","Cảm biến IoT · AIS · Camera · GPS"),
            ("LỚP SỐ", " · ".join(f["layers"])),
            ("AI + MÔ PHỎNG","Tối ưu · what-if · dự báo"),
            ("HÀNH ĐỘNG / KẾT QUẢ","Lệnh điều hành · cảnh báo · KPI")]
    fig,ax=plt.subplots(figsize=(11,2.7)); ax.axis("off"); ax.set_xlim(0,12); ax.set_ylim(0,2.6)
    cols=["#5B6472", col, "#7C5CFF", "#27C281"]; W=2.5; cx=[1.5,4.4,7.3,10.2]
    for i,(t,sub) in enumerate(stages):
        x=cx[i]
        ax.add_patch(FancyBboxPatch((x-W/2,0.5),W,1.5, boxstyle="round,pad=0.02,rounding_size=0.12",
                     linewidth=1.8, edgecolor=cols[i], facecolor=cols[i]+"18"))
        ax.text(x,1.62, wrp(t,18,2), fontsize=8.6, fontweight="bold", color=cols[i], ha="center", va="top")
        ax.text(x,1.02, wrp(sub,24,3), fontsize=6.8, color="#33415C", ha="center", va="top", linespacing=1.25)
        if i<3:
            ax.add_patch(FancyArrowPatch((x+W/2,1.25),(cx[i+1]-W/2,1.25), arrowstyle="-|>", mutation_scale=18, color="#9AA3B2", lw=1.8))
    save(fig,path); return path

def arch_img():
    path=os.path.join(IMG,"chatbot_arch.png")
    fig,ax=plt.subplots(figsize=(10.5,5.0)); ax.axis("off"); ax.set_xlim(0,12); ax.set_ylim(0,9)
    def box(cxx,cyy,w,h,t,sub,c):
        ax.add_patch(FancyBboxPatch((cxx-w/2,cyy-h/2),w,h,boxstyle="round,pad=0.02,rounding_size=0.12",lw=1.8,edgecolor=c,facecolor=c+"16"))
        ax.text(cxx,cyy+h/2-0.30,wrp(t,26,2),fontsize=9.2,fontweight="bold",color=c,ha="center",va="top")
        ax.text(cxx,cyy-0.05,wrp(sub,30,3),fontsize=6.8,color="#33415C",ha="center",va="top",linespacing=1.25)
    box(6,8.0,4.4,1.4,"Người dùng hỏi","ô chat tiếng Việt — câu hỏi tự nhiên","#0B1220")
    box(6,5.6,4.6,1.6,"copilot.js — Điều phối","lệnh thời gian · điều hướng · Q&A · hành động","#4D8DF6")
    box(2.3,2.4,3.5,1.9,"knowledge.js + faq.js","KB tĩnh ~180 mục · bỏ dấu · đồng nghĩa · IDF","#27C281")
    box(6.0,2.4,3.3,1.9,"locate.js","truy vấn thực thể sống · landmark · tổng hợp","#FF8A3C")
    box(9.7,2.4,3.5,1.9,"highlight · camera · cinematic","focus · highlight · tour điện ảnh · fork","#B07CFF")
    for (x0,y0,x1,y1) in [(6,7.3,6,6.4),(6,4.8,2.3,3.35),(6,4.8,6,3.35),(6,4.8,9.7,3.35)]:
        ax.add_patch(FancyArrowPatch((x0,y0),(x1,y1),arrowstyle="-|>",mutation_scale=15,color="#9AA3B2",lw=1.6))
    save(fig,path); return path

def timeline_img(f):
    col=f["color"]; path=os.path.join(IMG,f"timeline_{f['id']}.png")
    phases=[p[0] for p in f["plan"]]; n=len(phases)
    fig,ax=plt.subplots(figsize=(11,2.5)); ax.axis("off"); ax.set_xlim(0,12); ax.set_ylim(0,2.5)
    cols=[col,"#7C5CFF","#FF8A3C","#27C281"]; seg=11.4/n; x=0.3
    for i,ph in enumerate(phases):
        gd=ph.split("—")[0].strip() if "—" in ph else f"Giai đoạn {i+1}"
        name=ph.split("—")[-1].strip() if "—" in ph else ph
        c=cols[i%len(cols)]
        ax.add_patch(FancyBboxPatch((x+0.12,0.72),seg-0.44,0.95, boxstyle="round,pad=0.02,rounding_size=0.12",
                     lw=1.7, edgecolor=c, facecolor=c+"1A"))
        ax.text(x+seg/2-0.1,1.46,gd,fontsize=8.2,fontweight="bold",color=c,ha="center",va="top")
        ax.text(x+seg/2-0.1,1.10,wrp(name,20,2),fontsize=6.8,color="#33415C",ha="center",va="top",linespacing=1.2)
        if i<n-1:
            ax.add_patch(FancyArrowPatch((x+seg-0.30,1.18),(x+seg+0.12,1.18),arrowstyle="-|>",mutation_scale=15,color="#9AA3B2",lw=1.6))
        x+=seg
    ax.annotate("",xy=(11.8,0.42),xytext=(0.2,0.42),arrowprops=dict(arrowstyle="-|>",color="#9AA3B2",lw=1.4))
    ax.text(0.25,0.12,"Thời gian triển khai  →",fontsize=7,color="#8A93A6")
    save(fig,path); return path

def divergence_img():
    path=os.path.join(IMG,"divergence.png")
    xs=list(range(0,61))
    base=[3.0 for _ in xs]
    fork=[min(48,3.0+max(0,m-4)*0.78) for m in xs]
    mit=[min(20,3.0+max(0,m-4)*0.27) for m in xs]
    fig,ax=plt.subplots(figsize=(8.6,3.4))
    ax.plot(xs,base,color="#27C281",lw=2.2,label="Vũ trụ GỐC (baseline)")
    ax.plot(xs,fork,color="#E0556B",lw=2.6,label="Vũ trụ PHẢN THỰC (có sự cố)")
    ax.plot(xs,mit,color="#F8B23C",lw=2.2,ls="--",label="Sau khi AI hóa giải")
    ax.fill_between(xs,base,fork,color="#E0556B",alpha=0.10)
    ax.set_xlabel("phút sau sự cố",fontsize=9,color="#33415C")
    ax.set_ylabel("Chỉ số tác động tổng hợp",fontsize=9,color="#33415C")
    ax.tick_params(labelsize=8,colors="#5B6472")
    for s in ["top","right"]: ax.spines[s].set_visible(False)
    ax.spines["left"].set_color("#C8CED8"); ax.spines["bottom"].set_color("#C8CED8")
    ax.grid(axis="y",color="#EEF1F5",lw=0.8); ax.legend(fontsize=8,frameon=False,loc="upper left")
    save(fig,path); return path

def section(doc, n, title, accent):
    return h1(doc, f"{n}. {title}", accent)

# ── Tiện ích docx ──────────────────────────────────────────────────────────────
def hx(c): c=c.lstrip("#"); return RGBColor(int(c[0:2],16),int(c[2:4],16),int(c[4:6],16))

def add_cover(doc, title, subtitle, accent):
    logo=os.path.join(ASSETS,"logo_ndt.png")
    if os.path.exists(logo):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        try: p.add_run().add_picture(logo, width=Inches(1.4))
        except Exception: pass
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("DỰ ÁN BẢN SAO SỐ CẢNG BIỂN — NDT 15"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=hx("#5B6472")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=hx(accent)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle); r.font.size=Pt(12); r.font.color.rgb=hx("#33415C"); r.italic=True
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Báo cáo chi tiết · Tài liệu nội bộ"); r.font.size=Pt(10); r.font.color.rgb=hx("#8A93A6")
    doc.add_page_break()

def h1(doc, text, accent):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=hx(accent)
    return p

def body(doc, text):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(11); p.paragraph_format.space_after=Pt(6)
    return p

def bullets(doc, items):
    for it in items:
        p=doc.add_paragraph(style="List Bullet"); r=p.add_run(it); r.font.size=Pt(10.5)

def img(doc, path, caption, accent):
    if not os.path.exists(path): return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(6.2))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(caption); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=hx("#8A93A6")

def steplist(doc, steps):
    for i,(t,d) in enumerate(steps):
        p=doc.add_paragraph(); r=p.add_run(f"Bước {i+1} — {t}: "); r.bold=True; r.font.size=Pt(10.5)
        r2=p.add_run(d); r2.font.size=Pt(10.5); p.paragraph_format.space_after=Pt(4)

def scenario_table(doc, scenarios):
    tbl=doc.add_table(rows=1, cols=3)
    try: tbl.style="Table Grid"
    except Exception: pass
    hdr=tbl.rows[0].cells
    for c,txt in zip(hdr,["Kịch bản","Diễn biến (cascade)","Phương án hóa giải"]):
        c.paragraphs[0].add_run(txt).bold=True
    for nm,casc,mit in scenarios:
        c=tbl.add_row().cells; c[0].text=nm; c[1].text=casc; c[2].text=mit
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size=Pt(9)

def risklist(doc, risks):
    for rk,sol in risks:
        p=doc.add_paragraph(); r=p.add_run("⚠ Rủi ro: "); r.bold=True; r.font.size=Pt(10.5)
        r2=p.add_run(rk); r2.font.size=Pt(10.5)
        p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Inches(0.3)
        r3=p2.add_run("→ Giải pháp: "); r3.bold=True; r3.font.size=Pt(10); r3.font.color.rgb=hx("#1F9D6B")
        r4=p2.add_run(sol); r4.font.size=Pt(10); p2.paragraph_format.space_after=Pt(6)

def build_feature(f):
    doc=Document(); acc=f["color"]
    add_cover(doc, f"TÍNH NĂNG {f['id']} — {f['name'].upper()}", f["desc"], acc)
    n=1
    section(doc,n,"Bối cảnh & tầm quan trọng",acc); n+=1; body(doc,f.get("context","")); body(doc,f["intro"])
    section(doc,n,"Vấn đề thực tế (Pain Points)",acc); n+=1; body(doc,f["pain"])
    section(doc,n,"Giải pháp & mô tả tính năng",acc); n+=1
    body(doc,"Giải pháp tổng quan: "+f["desc"]); body(doc,f["mech"])
    section(doc,n,"Cơ chế hoạt động & quy trình",acc); n+=1
    steplist(doc,f["steps"]); img(doc, flow_img(f), f"Sơ đồ {len(f['steps'])} bước vận hành — {f['short']}", acc)
    section(doc,n,"Phân tích chuyên sâu quy trình",acc); n+=1; body(doc,f.get("deepdive",""))
    section(doc,n,"Kiến trúc dữ liệu & công nghệ",acc); n+=1
    body(doc,f.get("tech","")); img(doc, pipeline_img(f), "Pipeline Digital Twin: Nguồn dữ liệu → Lớp số → AI/Mô phỏng → Hành động", acc)
    section(doc,n,"Digital Twin được ứng dụng như thế nào",acc); n+=1; body(doc,f["dt_apply"])
    if f.get("scenarios"):
        img(doc, divergence_img(), "Minh họa: chỉ số tác động phân kỳ giữa Gốc và Phản thực (và sau khi AI hóa giải)", acc)
        section(doc,n,"Các kịch bản mô phỏng",acc); n+=1; scenario_table(doc, f["scenarios"])
    section(doc,n,"Ứng dụng thực tế",acc); n+=1; bullets(doc,f["apps"])
    section(doc,n,"Chỉ số đo lường (KPI) & lợi ích",acc); n+=1
    img(doc, kpi_img(f), "Các chỉ số (KPI) tiêu biểu", acc)
    body(doc,"Lợi ích kinh tế (ROI): "+f.get("roi","")); bullets(doc,f["benefits"])
    section(doc,n,"Rủi ro & giải pháp",acc); n+=1; risklist(doc,f.get("risks",[]))
    section(doc,n,"Lộ trình triển khai",acc); n+=1
    for ph,ct in f["plan"]:
        p=doc.add_paragraph(); r=p.add_run(ph+": "); r.bold=True; r.font.size=Pt(10.5); r2=p.add_run(ct); r2.font.size=Pt(10.5)
    img(doc, timeline_img(f), "Lộ trình triển khai theo giai đoạn", acc)
    section(doc,n,"Kết luận & khuyến nghị",acc); n+=1; body(doc,f.get("conclusion",""))
    out=os.path.join(HERE, f"TinhNang_{f['id']}_{f['short'].replace(' ','')}.docx")
    doc.save(out); return out

def coverage_img():
    path=os.path.join(IMG,"coverage.png")
    cats=["Truy vấn động (bãi/tàu/xe…)","Vật thể & giới thiệu","Vận hành cảng","Digital Twin & 15 lớp",
          "Công nghệ (IoT/AI/BIM…)","ESG & Năng lượng","An ninh & Chống chịu","Khách hàng/Thương mại"]
    vals=[2500,1200,1100,900,800,700,600,500]
    cols=["#B07CFF","#34E0F0","#27C281","#4D8DF6","#7C5CFF","#2ADA9A","#E0556B","#FF8A3C"]
    fig,ax=plt.subplots(figsize=(8.8,3.8)); y=list(range(len(cats)))
    ax.barh(y,vals,color=cols,height=0.62); ax.invert_yaxis()
    ax.set_yticks(y); ax.set_yticklabels(cats,fontsize=8.5,color="#33415C")
    for i,v in enumerate(vals): ax.text(v+40,i,f"~{v}",va="center",fontsize=8,color="#5B6472")
    ax.set_xlabel("Số cách hỏi ước tính",fontsize=9,color="#33415C"); ax.set_xlim(0,2900)
    for s in ["top","right"]: ax.spines[s].set_visible(False)
    ax.spines["left"].set_color("#C8CED8"); ax.spines["bottom"].set_color("#C8CED8")
    ax.tick_params(labelsize=8,colors="#5B6472")
    save(fig,path); return path

def example_table(doc, rows):
    tbl=doc.add_table(rows=1, cols=2)
    try: tbl.style="Table Grid"
    except Exception: pass
    h=tbl.rows[0].cells
    for c,t in zip(h,["Câu hỏi mẫu của khách","Hệ xử lý như thế nào"]): c.paragraphs[0].add_run(t).bold=True
    for q,a in rows:
        c=tbl.add_row().cells; c[0].text=q; c[1].text=a
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size=Pt(9)

def build_chatbot():
    acc="#4D8DF6"; doc=Document()
    add_cover(doc,"XÂY DỰNG TRỢ LÝ ẢO (CHATBOT) NDT 15","Hệ hỏi-đáp & điều khiển bản sao số bằng ngôn ngữ tự nhiên", acc)
    n=1
    section(doc,n,"Giới thiệu & mục tiêu",acc); n+=1
    body(doc,"Trợ lý ảo (Copilot) là 'người dẫn đường' cho bản sao số NDT15: trả lời mọi câu hỏi về cảng bằng tiếng Việt tự nhiên và TRỰC TIẾP điều khiển cảnh 3D (bay camera, highlight vật thể, chạy mô phỏng, dẫn tham quan). Mục tiêu là đáp ứng 5.000–10.000 câu hỏi thường gặp mà không cần backend hay dịch vụ LLM bên ngoài — chạy hoàn toàn trên trình duyệt.")
    section(doc,n,"Bối cảnh & vì sao cần chatbot",acc); n+=1
    body(doc,"Một bản sao số giàu dữ liệu sẽ vô dụng nếu người dùng không biết cách khai thác. Khách tham quan, lãnh đạo và đối tác không muốn học thao tác phức tạp — họ muốn 'hỏi là ra'. Trợ lý ảo hạ rào cản sử dụng xuống bằng 0: người dùng hỏi bằng lời như nói chuyện, hệ thống vừa trả lời vừa tự lái camera tới đúng nơi, biến twin từ công cụ chuyên gia thành trải nghiệm cho mọi người.")
    section(doc,n,"Kiến trúc tổng thể",acc); n+=1
    body(doc,"Hệ gồm các module phối hợp: copilot.js điều phối tổng (nhận câu hỏi, phân loại lệnh/điều hướng/Q&A); knowledge.js + faq.js là ngân hàng tri thức tĩnh (~180 mục); locate.js truy vấn thực thể đang chạy; còn highlight.js / camera.js / cinematic.js thực thi hành động trên cảnh 3D.")
    img(doc, arch_img(), "Kiến trúc trợ lý ảo NDT15", acc)
    section(doc,n,"Cơ chế truy hồi tri thức",acc); n+=1
    body(doc,"Câu hỏi được chuẩn hóa qua 3 bước: (1) bỏ dấu tiếng Việt (nên 'bãi'='bai', không phụ thuộc gõ Telex/VNI); (2) gộp từ đồng nghĩa (tàu=thuyền=ship=vessel, cẩu=crane…); (3) tách token và loại từ dừng. Sau đó hệ chấm điểm theo trọng số IDF — từ khóa càng đặc trưng càng quyết định — kèm khớp một phần cho biến thể hình thái.")
    body(doc,"Hai tinh chỉnh quan trọng đảm bảo độ chính xác: các từ 'đệm' (thế nào, ở đâu, hoạt động…) bị hạ trọng số để từ khóa THỰC THỂ chiếm ưu thế; và hệ yêu cầu có ít nhất một token khớp dài ≥3 ký tự — nhờ vậy câu vô nghĩa sẽ trả về gợi ý thay vì đoán bừa. Mỗi mục trả lời do đó khớp được hàng chục cách hỏi khác nhau.")
    section(doc,n,"Truy vấn thực thể sống & hành động",acc); n+=1
    body(doc,"Với câu hỏi về đối tượng cụ thể (bãi số 15, tàu OCEAN, xe ở bãi 3, tàu hỏa, tổng container tồn kho, nhà máy hydro…), locate.js đọc trạng thái real-time từ scene rồi HIGHLIGHT đối tượng + FOCUS camera + trả thông tin. Nếu đang ở tầng hầm mà cần ra mặt đất, hệ tự 'nổi lên' trước khi bay. Các hành động chính:")
    bullets(doc,["Điều hướng: 'dẫn tôi đến cổng/bãi/nhà máy hydro…' — bay camera (tự nổi lên mặt đất nếu đang dưới hầm).",
                 "Tham quan điện ảnh: 'dẫn tôi đi tham quan' — camera xoay quanh & zoom ra/vào từng khu, kèm thuyết minh.",
                 "Chạy tính năng: 'chạy mô phỏng tính năng bến' — mở mô phỏng Digital Twin (phân biệt với 'là gì/ở đâu' chỉ chọn).",
                 "Mô phỏng sự cố: 'nếu bão thì sao' — phân nhánh thực tại (fork) và mô phỏng hậu quả.",
                 "Tua thời gian: 'tua lại', 'cho xem tương lai', 'về thực tại'.",
                 "Tổng hợp: 'có bao nhiêu container tồn kho', 'xe tải nào đang làm gì'."])
    section(doc,n,"Ví dụ câu hỏi & cách xử lý",acc); n+=1
    example_table(doc,[
        ("Cẩu RTG là gì?","Khớp KB tĩnh → giải thích + chọn tính năng liên quan."),
        ("Bãi container 15 nằm ở đâu?","Truy vấn động → highlight bãi 15 + bay camera + thông tin."),
        ("Tàu OCEAN đang làm gì?","Đọc trạng thái tàu real-time → bám camera + mô tả."),
        ("Có bao nhiêu container tồn kho?","Tổng hợp từ 30 bãi → trả số liệu + focus cụm bãi."),
        ("Dẫn tôi đến nhà máy hydro","Điều hướng landmark → nổi lên mặt đất + bay tới + highlight."),
        ("Chạy mô phỏng tính năng cổng","Mở mô phỏng Digital Twin của tính năng cổng."),
        ("Nếu bão thì sao?","Fork kịch bản bão+ngập → cascade + biểu đồ + hóa giải."),
        ("Dẫn tôi đi tham quan","Tour điện ảnh: orbit + zoom qua từng khu, có thuyết minh."),
    ])
    section(doc,n,"Độ phủ câu hỏi",acc); n+=1
    body(doc,"Khoảng 180 mục trả lời tĩnh × hàng chục cách hỏi mỗi mục (~3.000–5.000 cách hỏi), cộng với engine động (30 bãi × nhiều loại hỏi, tàu/xe/tàu hỏa/AGV, 11 landmark, tổng hợp container…) → phủ thoải mái 5.000–10.000 câu hỏi thực tế.")
    img(doc, coverage_img(), "Ước lượng độ phủ câu hỏi theo nhóm chủ đề", acc)
    section(doc,n,"Lợi ích",acc); n+=1
    bullets(doc,["Hạ rào cản sử dụng twin về 0 — ai cũng 'hỏi là ra'.",
                 "Vừa trả lời vừa hành động (lái camera, highlight, mô phỏng) → trực quan, thuyết phục.",
                 "Chạy offline trên trình duyệt, không phụ thuộc dịch vụ ngoài, bảo mật & chi phí thấp.",
                 "Công cụ trình diễn mạnh khi giới thiệu với lãnh đạo, nhà đầu tư, khách tham quan."])
    section(doc,n,"Rủi ro & giải pháp",acc); n+=1
    risklist(doc,[("Câu hỏi quá phức tạp/đa ý","Cho phép ghép LLM phía sau cùng giao diện parse; tách câu thành ý nhỏ."),
                  ("Thiếu dữ liệu chuyên biệt (giá, thủ tục)","Bổ sung vào FAQ — chỉ thêm dữ liệu, không sửa code."),
                  ("Hiểu sai do biến thể ngôn ngữ","Mở rộng từ điển đồng nghĩa & ngưỡng tin cậy; ghi log câu trượt để cải thiện.")])
    section(doc,n,"Lộ trình mở rộng",acc); n+=1
    cb=dict(id="CB", color=acc, plan=[
        ("Giai đoạn 1 — Ngân hàng tri thức","Catalog + FAQ tĩnh, cơ chế bỏ dấu/đồng nghĩa/IDF."),
        ("Giai đoạn 2 — Truy vấn động","locate.js đọc thực thể sống + highlight/focus + tổng hợp."),
        ("Giai đoạn 3 — Hành động & tham quan","Bay camera, tour điện ảnh, mô phỏng/fork, tua thời gian."),
        ("Giai đoạn 4 — Mở rộng","FAQ đặc thù, hỗ trợ giọng nói, tùy chọn ghép LLM.")])
    for ph,ct in cb["plan"]:
        p=doc.add_paragraph(); r=p.add_run(ph+": "); r.bold=True; r.font.size=Pt(10.5); r2=p.add_run(ct); r2.font.size=Pt(10.5)
    img(doc, timeline_img(cb), "Lộ trình xây dựng & mở rộng chatbot", acc)
    section(doc,n,"Kết luận",acc); n+=1
    body(doc,"Trợ lý ảo NDT15 biến một bản sao số phức tạp thành trải nghiệm 'hỏi là ra' cho mọi đối tượng, vừa trả lời vừa hành động trong cảnh 3D. Đây là cầu nối giữa sức mạnh dữ liệu của twin và người dùng cuối — nên được xem là thành phần lõi, không phải tiện ích phụ.")
    out=os.path.join(HERE,"BaoCao_XayDung_Chatbot.docx"); doc.save(out); return out

if __name__=="__main__":
    outs=[]
    for f in FEATURES: outs.append(build_feature(f))
    outs.append(build_chatbot())
    print("ĐÃ TẠO", len(outs), "FILE:")
    for o in outs: print(" -", os.path.basename(o))
